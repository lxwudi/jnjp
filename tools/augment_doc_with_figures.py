from __future__ import annotations

import math
import textwrap
from pathlib import Path

from PIL import Image, ImageColor, ImageDraw, ImageFilter, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path("/Users/flx/Documents/jienengjianpai")
INPUT_DOC = ROOT / "校园交换机智能体节能平台设计与实现.docx"
OUTPUT_DOC = ROOT / "校园交换机智能体节能平台设计与实现-加图版.docx"
SHOT_DIR = Path("/tmp/doc-ui-shots/png")
ASSET_DIR = Path("/tmp/jnjp-doc-figures")
ASSET_DIR.mkdir(parents=True, exist_ok=True)

FONT_SONGTI = "/System/Library/Fonts/Supplemental/Songti.ttc"
FONT_SANS = FONT_SONGTI
FONT_SERIF = FONT_SONGTI

BG = "#F6FBFC"
CARD = "#FFFFFF"
TEXT = "#253847"
MUTED = "#708896"
TEAL = "#2E5A64"
TEAL_2 = "#5AAFC0"
MINT = "#9BD47F"
SKY = "#8ED2EA"
LIME = "#E5F08A"
BORDER = "#D6E4E8"
PALE = "#EDF6F7"
SHADOW = "#CFE0E5"
GOLD = "#F0B44D"


def load_font(path: str, size: int) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(path, size)


def text_size(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont) -> tuple[int, int]:
    bbox = draw.textbbox((0, 0), text, font=font)
    return bbox[2] - bbox[0], bbox[3] - bbox[1]


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for ch in text:
        test = current + ch
        if text_size(draw, test, font)[0] <= max_width or not current:
            current = test
        else:
            lines.append(current)
            current = ch
    if current:
        lines.append(current)
    return lines


def draw_multiline(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int],
    text: str,
    font: ImageFont.FreeTypeFont,
    fill: str,
    max_width: int,
    line_gap: int = 8,
) -> int:
    x, y = xy
    total_h = 0
    for paragraph in text.split("\n"):
        lines = wrap_text(draw, paragraph, font, max_width) if paragraph else [""]
        for line in lines:
            draw.text((x, y + total_h), line, font=font, fill=fill)
            total_h += text_size(draw, line or "国", font)[1] + line_gap
        total_h += 6
    return total_h


def draw_text_center(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font: ImageFont.FreeTypeFont,
    fill: str,
):
    tw, th = text_size(draw, text, font)
    x = box[0] + (box[2] - box[0] - tw) / 2
    y = box[1] + (box[3] - box[1] - th) / 2 - 2
    draw.text((x, y), text, font=font, fill=fill)


def rounded_card(
    base: Image.Image,
    box: tuple[int, int, int, int],
    radius: int = 36,
    fill: str = CARD,
    outline: str | None = None,
    shadow_offset: tuple[int, int] = (0, 10),
):
    x1, y1, x2, y2 = box
    shadow = Image.new("RGBA", base.size, (0, 0, 0, 0))
    shadow_draw = ImageDraw.Draw(shadow)
    sx, sy = shadow_offset
    shadow_draw.rounded_rectangle((x1 + sx, y1 + sy, x2 + sx, y2 + sy), radius=radius, fill=ImageColor.getrgb(SHADOW) + (90,))
    shadow = shadow.filter(ImageFilter.GaussianBlur(18))
    base.alpha_composite(shadow)
    draw = ImageDraw.Draw(base)
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=2 if outline else 0)


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str, width: int = 6):
    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    head_len = 20
    left = (
        end[0] - head_len * math.cos(angle - math.pi / 6),
        end[1] - head_len * math.sin(angle - math.pi / 6),
    )
    right = (
        end[0] - head_len * math.cos(angle + math.pi / 6),
        end[1] - head_len * math.sin(angle + math.pi / 6),
    )
    draw.polygon([end, left, right], fill=color)


def draw_pill(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], fill: str, text: str, font, text_fill: str = "white"):
    draw.rounded_rectangle(box, radius=(box[3] - box[1]) // 2, fill=fill)
    tw, th = text_size(draw, text, font)
    x = box[0] + (box[2] - box[0] - tw) / 2
    y = box[1] + (box[3] - box[1] - th) / 2 - 2
    draw.text((x, y), text, font=font, fill=text_fill)


def create_energy_chart() -> Path:
    img = Image.new("RGBA", (1800, 980), BG)
    rounded_card(img, (70, 70, 1730, 910), radius=44, fill=CARD, outline="#E4EEF0")
    draw = ImageDraw.Draw(img)
    sans_36 = load_font(FONT_SANS, 36)
    sans_30 = load_font(FONT_SANS, 30)
    sans_26 = load_font(FONT_SANS, 26)
    sans_24 = load_font(FONT_SANS, 24)
    sans_22 = load_font(FONT_SANS, 22)
    sans_20 = load_font(FONT_SANS, 20)
    serif_22 = load_font(FONT_SERIF, 22)
    serif_20 = load_font(FONT_SERIF, 20)
    serif_18 = load_font(FONT_SERIF, 18)

    draw.text((130, 115), "背景数据支撑", font=sans_24, fill=TEAL_2)
    draw.text((130, 160), "能耗构成与改造重点", font=sans_36, fill=TEXT)
    draw.text((130, 214), "根据正文中的背景数据，将校园机房负荷、交换机占比与接入层改造重点进行可视化归纳。", font=serif_22, fill=MUTED)

    rounded_card(img, (110, 285, 980, 850), radius=34, fill="#FBFEFE", outline="#E3ECEF")
    rounded_card(img, (1020, 285, 1670, 850), radius=34, fill="#FBFEFE", outline="#E3ECEF")

    draw.text((155, 320), "分层拆解路径", font=sans_30, fill=TEXT)
    draw.text((155, 364), "从综合负荷逐层收敛到端口级治理对象，直接解释“为什么先做接入层节能”。", font=serif_20, fill=MUTED)

    bars = [
        {
            "label": "综合负荷",
            "x": 220,
            "y": 430,
            "w": 620,
            "h": 56,
            "ratio": 0.44,
            "fill": TEAL,
            "rest": "#E8F0F2",
            "left_text": "IT 设备 44%",
            "right_text": "非 IT 56%",
        },
        {
            "label": "IT 设备内部",
            "x": 290,
            "y": 545,
            "w": 550,
            "h": 56,
            "ratio": 0.30,
            "fill": SKY,
            "rest": "#EDF6F8",
            "left_text": "交换机 30%",
            "right_text": "其他 IT 设备 70%",
        },
        {
            "label": "交换机内部",
            "x": 360,
            "y": 660,
            "w": 480,
            "h": 56,
            "ratio": 0.85,
            "fill": MINT,
            "rest": "#F1F7EE",
            "left_text": "接入层 >85%",
            "right_text": "汇聚/核心 <15%",
        },
    ]

    for idx, bar in enumerate(bars, start=1):
        badge_x = bar["x"] - 70
        badge_y = bar["y"] + 10
        draw.rounded_rectangle((badge_x, badge_y, badge_x + 42, badge_y + 36), radius=18, fill="#EAF5F6")
        draw_text_center(draw, (badge_x, badge_y, badge_x + 42, badge_y + 36), str(idx), sans_22, TEAL)
        draw.text((bar["x"], bar["y"] - 42), bar["label"], font=sans_26, fill=TEXT)

        draw.rounded_rectangle((bar["x"], bar["y"], bar["x"] + bar["w"], bar["y"] + bar["h"]), radius=28, fill=bar["rest"])
        fill_w = int(bar["w"] * bar["ratio"])
        draw.rounded_rectangle((bar["x"], bar["y"], bar["x"] + fill_w, bar["y"] + bar["h"]), radius=28, fill=bar["fill"])

        left_box = (bar["x"], bar["y"], bar["x"] + fill_w, bar["y"] + bar["h"])
        draw_text_center(draw, left_box, bar["left_text"], sans_22, "#11313D")

        right_area_w = bar["w"] - fill_w
        right_box = (bar["x"] + fill_w, bar["y"], bar["x"] + bar["w"], bar["y"] + bar["h"])
        draw_text_center(draw, right_box, bar["right_text"], sans_20, MUTED)

    arrow(draw, (530, 486), (565, 535), TEAL_2, width=6)
    arrow(draw, (585, 601), (620, 650), "#8CBF80", width=6)

    rounded_card(img, (175, 748, 900, 830), radius=26, fill="#F7FBF4")
    draw.line((430, 762, 430, 816), fill="#DCE8DE", width=2)
    draw.line((640, 762, 640, 816), fill="#DCE8DE", width=2)
    draw_text_center(draw, (195, 760, 420, 812), "44% × 30% × 85% ≈", sans_22, MUTED)
    draw_text_center(draw, (440, 756, 630, 816), "11.2%+", sans_36, TEAL)
    draw_multiline(draw, (662, 768), "综合负荷中可优先纳入\n端口级节能治理的部分", serif_18, TEXT, 208, line_gap=2)

    draw.text((1060, 320), "治理启示", font=sans_30, fill=TEXT)
    draw.text((1060, 364), "把数据结论直接翻译成治理顺序、执行边界和闭环能力。", font=serif_20, fill=MUTED)

    insight_cards = [
        ("先抓接入层", "接入层交换机在交换机能耗中占比最高，最适合作为第一阶段的节能切入口。", "#EEF7F7", SKY),
        ("先易后难", "首批对象以空闲端口、低负载端口和低风险端口为主，优先保证收益与稳定性。", "#F3F8EE", MINT),
        ("纳入闭环", "智能识别 -> 护栏校验 -> 自动执行 -> 审计追踪，让节能过程可控、可查、可复盘。", "#FFF8EC", GOLD),
    ]
    card_top = 420
    for idx, (title, body, fill, accent) in enumerate(insight_cards, start=1):
        y = card_top + (idx - 1) * 135
        rounded_card(img, (1055, y, 1625, y + 114), radius=26, fill=fill)
        draw.rounded_rectangle((1080, y + 22, 1120, y + 58), radius=18, fill=accent)
        draw_text_center(draw, (1080, y + 22, 1120, y + 58), str(idx), sans_22, "#15323D")
        draw.text((1146, y + 16), title, font=sans_26, fill=TEXT)
        draw_multiline(draw, (1146, y + 52), body, serif_20, MUTED, 426, line_gap=4)

    path = ASSET_DIR / "energy-composition.png"
    img.convert("RGB").save(path, quality=95)
    return path


def create_architecture_diagram() -> Path:
    img = Image.new("RGBA", (1800, 900), BG)
    rounded_card(img, (70, 70, 1730, 830), radius=44, fill=CARD, outline="#E4EEF0")
    draw = ImageDraw.Draw(img)
    sans_34 = load_font(FONT_SANS, 34)
    sans_30 = load_font(FONT_SANS, 30)
    sans_24 = load_font(FONT_SANS, 24)
    sans_22 = load_font(FONT_SANS, 22)
    serif_22 = load_font(FONT_SERIF, 22)
    serif_20 = load_font(FONT_SERIF, 20)

    draw.text((130, 118), "前后端与数据持久化协同关系", font=sans_34, fill=TEXT)
    draw.text((130, 165), "围绕“前端展示—服务编排—数据沉淀”三层结构整理平台实现路径，并显式标出知识库与规则兜底。", font=serif_22, fill=MUTED)

    layer_boxes = [
        ((130, 280, 1670, 400), "#F5FBFC", "前端展示层", [("总览驾驶舱", TEAL_2), ("自治主控", TEAL_2), ("统计分析", TEAL_2), ("审计追踪", TEAL_2)]),
        ((130, 445, 1670, 585), "#F6FAF5", "服务编排层", [("接口管理", MINT), ("自治调度", MINT), ("智能体规划", MINT), ("执行引擎", MINT), ("日志服务", MINT)]),
        ((130, 630, 1670, 760), "#FBFBF5", "数据与知识层", [("console.db", GOLD), ("agent-runtime.db", GOLD), ("RAG 知识库", GOLD)]),
    ]

    for box, fill, title, chips in layer_boxes:
        rounded_card(img, box, radius=32, fill=fill, outline="#DDE8EA")
        draw.text((box[0] + 28, box[1] + 22), title, font=sans_30, fill=TEXT)
        chip_x = box[0] + 280
        chip_y = box[1] + 22
        for label, color in chips:
            width = text_size(draw, label, sans_24)[0] + 48
            draw.rounded_rectangle((chip_x, chip_y, chip_x + width, chip_y + 48), radius=24, fill=color)
            draw.text((chip_x + 24, chip_y + 10), label, font=sans_24, fill="#13313D")
            chip_x += width + 20

    arrow(draw, (900, 402), (900, 440), TEAL_2, width=8)
    arrow(draw, (900, 587), (900, 625), MINT, width=8)

    decision_box = (1268, 88, 1648, 248)
    rounded_card(img, decision_box, radius=28, fill="#F1F7F8", outline="#DDE8EA", shadow_offset=(0, 5))
    draw.text((1298, 111), "双模式决策", font=sans_22, fill=TEAL)
    draw_pill(draw, (1298, 147, 1468, 191), TEAL, "模型智能体", sans_24)
    draw_pill(draw, (1484, 147, 1618, 191), "#6D8792", "规则兜底", sans_24)
    draw_multiline(
        draw,
        (1298, 205),
        "统一汇入服务编排层，保障演示与自治连续可用。",
        serif_20,
        MUTED,
        300,
        line_gap=1,
    )

    path = ASSET_DIR / "architecture-diagram.png"
    img.convert("RGB").save(path, quality=95)
    return path


def create_flow_diagram() -> Path:
    img = Image.new("RGBA", (1800, 980), BG)
    rounded_card(img, (70, 70, 1730, 910), radius=44, fill=CARD, outline="#E4EEF0")
    draw = ImageDraw.Draw(img)
    sans_34 = load_font(FONT_SANS, 34)
    sans_28 = load_font(FONT_SANS, 28)
    sans_22 = load_font(FONT_SANS, 22)
    serif_20 = load_font(FONT_SERIF, 20)

    draw.text((130, 118), "自治闭环、RAG 检索与护栏约束协同流程", font=sans_34, fill=TEXT)
    draw.text((130, 164), "以六步闭环为主线，体现数据、知识、边界控制与执行留痕之间的顺序关系。", font=serif_20, fill=MUTED)

    boxes = {
        "1": (120, 300, 540, 455),
        "2": (690, 300, 1110, 455),
        "3": (1260, 300, 1680, 455),
        "4": (1260, 560, 1680, 715),
        "5": (690, 560, 1110, 715),
        "6": (120, 560, 540, 715),
    }
    items = {
        "1": ("接入端口数据", "汇聚接口信息，形成可治理接口池。", TEAL_2),
        "2": ("检索知识依据", "从设备手册、制度与复盘中提取匹配策略。", SKY),
        "3": ("设置执行边界", "利用率阈值、连接数阈值、时窗与动作边界协同生效。", MINT),
        "4": ("智能体自动巡检", "结合实时状态与知识依据评估风险与收益。", TEAL_2),
        "5": ("自动执行 / 跳过", "低风险动作自动放行，高风险动作保留并说明原因。", GOLD),
        "6": ("结果展示与审计", "回写收益、输出日志，并为下一轮优化提供依据。", "#B8E06C"),
    }

    for key, box in boxes.items():
        title, body, accent = items[key]
        rounded_card(img, box, radius=28, fill="#FAFCFC", outline="#DDE8EA")
        draw.rounded_rectangle((box[0] + 22, box[1] + 22, box[0] + 72, box[1] + 72), radius=16, fill=accent)
        draw.text((box[0] + 36, box[1] + 30), key, font=sans_28, fill="#12303C")
        draw.text((box[0] + 95, box[1] + 24), title, font=sans_28, fill=TEXT)
        draw_multiline(draw, (box[0] + 95, box[1] + 72), body, serif_20, MUTED, box[2] - box[0] - 125, line_gap=5)

    draw_pill(draw, (900, 250, 980, 288), TEAL, "RAG", sans_22)
    draw_pill(draw, (1448, 250, 1568, 288), "#6FAE8A", "护栏", sans_22)

    arrow(draw, (540, 378), (680, 378), SKY, 8)
    arrow(draw, (1110, 378), (1250, 378), MINT, 8)
    arrow(draw, (1470, 455), (1470, 548), MINT, 8)
    arrow(draw, (1260, 638), (1120, 638), GOLD, 8)
    arrow(draw, (690, 638), (550, 638), "#B6D96E", 8)
    arrow(draw, (300, 552), (300, 465), TEAL_2, 8)

    path = ASSET_DIR / "flow-diagram.png"
    img.convert("RGB").save(path, quality=95)
    return path


def add_shadow_panel(target: Image.Image, box: tuple[int, int, int, int], fill: str = "#FFFFFF"):
    rounded_card(target, box, radius=24, fill=fill, outline="#DCE7EA", shadow_offset=(0, 8))


def crop_ui_shot(name: str, box: tuple[int, int, int, int]) -> Image.Image:
    img = Image.open(SHOT_DIR / f"{name}.png").convert("RGB")
    return img.crop(box)


def create_ui_gallery() -> Path:
    img = Image.new("RGBA", (1800, 1240), BG)
    rounded_card(img, (50, 50, 1750, 1180), radius=44, fill="#FCFEFE", outline="#E4EEF0")
    draw = ImageDraw.Draw(img)
    sans_34 = load_font(FONT_SANS, 34)
    sans_24 = load_font(FONT_SANS, 24)
    sans_20 = load_font(FONT_SANS, 20)

    draw.text((110, 95), "前端核心页面组合展示", font=sans_34, fill=TEXT)
    draw.text((110, 142), "围绕总览、自治、护栏、统计与审计五个页面截取核心区域，便于在说明书中集中展示系统界面。", font=sans_20, fill=MUTED)

    crop_boxes = {
        "overview-mid-visible": (14, 80, 476, 760),
        "agent-overview-visible": (14, 80, 476, 760),
        "agent-guardrails-visible": (14, 80, 476, 760),
        "insight-mid-visible": (14, 10, 476, 760),
        "audit-visible": (14, 10, 476, 760),
    }
    panels = [
        ("图1 总览驾驶舱", "overview-mid-visible"),
        ("图2 自治主控台", "agent-overview-visible"),
        ("图3 执行边界", "agent-guardrails-visible"),
        ("图4 统计分析", "insight-mid-visible"),
        ("图5 审计追踪", "audit-visible"),
    ]
    positions = [
        (105, 245),
        (650, 245),
        (1195, 245),
        (375, 720),
        (920, 720),
    ]

    for (label, key), (x, y) in zip(panels, positions):
        panel_w, panel_h = 500, 410
        add_shadow_panel(img, (x, y, x + panel_w, y + panel_h), fill="#FFFFFF")
        draw.rounded_rectangle((x + 22, y + 18, x + 170, y + 54), radius=18, fill="#EAF5F6")
        draw.text((x + 36, y + 24), label, font=sans_20, fill=TEXT)
        shot = crop_ui_shot(key, crop_boxes[key]).resize((454, 300), Image.Resampling.LANCZOS)
        img.paste(shot, (x + 23, y + 84))
        draw.rounded_rectangle((x + 22, y + 84, x + 476, y + 384), radius=20, outline="#D6E4E8", width=2)
        draw.text((x + 28, y + 392), "界面实拍", font=sans_20, fill=MUTED)

    path = ASSET_DIR / "ui-gallery.png"
    img.convert("RGB").save(path, quality=95)
    return path


def set_run_font(run, font_name: str, size_pt: float, bold: bool = False, color: str | None = None):
    run.font.name = font_name
    run.bold = bold
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii", font_name)
    rfonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi", font_name)
    rfonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", font_name)


def insert_paragraph_after(paragraph: Paragraph, text: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def add_picture_after(paragraph: Paragraph, image_path: Path, width_cm: float) -> Paragraph:
    img_para = insert_paragraph_after(paragraph)
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.paragraph_format.space_before = Pt(6)
    img_para.paragraph_format.space_after = Pt(3)
    run = img_para.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    return img_para


def add_caption_after(paragraph: Paragraph, text: str) -> Paragraph:
    cap_para = insert_paragraph_after(paragraph)
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_para.paragraph_format.space_before = Pt(2)
    cap_para.paragraph_format.space_after = Pt(8)
    run = cap_para.add_run(text)
    set_run_font(run, "宋体", 10.5)
    return cap_para


def build_doc():
    energy = create_energy_chart()
    arch = create_architecture_diagram()
    flow = create_flow_diagram()
    gallery = create_ui_gallery()

    doc = Document(INPUT_DOC)
    paras = doc.paragraphs

    ref = paras[11]
    p = add_picture_after(ref, energy, 15.2)
    add_caption_after(p, "校园机房能耗构成与接入层改造重点示意")

    ref = paras[39]
    p = add_picture_after(ref, arch, 15.2)
    add_caption_after(p, "平台前后端、自治服务与数据持久化协同关系示意")

    ref = paras[49]
    p = add_picture_after(ref, flow, 15.2)
    add_caption_after(p, "自治闭环、知识检索与执行边界协同流程示意")

    ref = paras[61]
    p = add_picture_after(ref, gallery, 15.4)
    add_caption_after(p, "系统五大核心页面组合展示")

    doc.save(OUTPUT_DOC)
    print(OUTPUT_DOC)


if __name__ == "__main__":
    build_doc()
