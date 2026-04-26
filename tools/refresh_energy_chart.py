from __future__ import annotations

import importlib.util
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


ROOT = Path("/Users/flx/Documents/jienengjianpai")
DOC_PATHS = [
    ROOT / "校园交换机智能体节能平台设计与实现.docx",
    ROOT / "校园交换机智能体节能平台设计与实现-加图版.docx",
]
MODULE_PATH = ROOT / "tools" / "augment_doc_with_figures.py"


def load_augment_module():
    spec = importlib.util.spec_from_file_location("augment_doc_with_figures", MODULE_PATH)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


def paragraph_has_drawing(paragraph: Paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing"))


def remove_paragraph(paragraph: Paragraph):
    parent = paragraph._element.getparent()
    if parent is not None:
        parent.remove(paragraph._element)


def insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def add_picture_after(paragraph: Paragraph, image_path: Path, width_cm: float) -> Paragraph:
    img_para = insert_paragraph_after(paragraph)
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.paragraph_format.space_before = Pt(6)
    img_para.paragraph_format.space_after = Pt(3)
    run = img_para.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    return img_para


def refresh_doc(doc_path: Path, caption: str, image_path: Path):
    doc = Document(doc_path)
    paragraphs = list(doc.paragraphs)
    for i, paragraph in enumerate(paragraphs):
        if paragraph.text != caption:
            continue

        anchor = None
        for j in range(i - 1, -1, -1):
            if paragraphs[j].text.strip():
                anchor = paragraphs[j]
                break
        if anchor is None:
            break

        for j in range(i - 1, -1, -1):
            if paragraphs[j] == anchor:
                break
            if paragraph_has_drawing(paragraphs[j]) or not paragraphs[j].text.strip():
                remove_paragraph(paragraphs[j])

        add_picture_after(anchor, image_path, 15.2)
        break

    doc.save(doc_path)


def main():
    module = load_augment_module()
    assets = {
        "校园机房能耗构成与接入层改造重点示意": module.create_energy_chart(),
        "平台前后端、自治服务与数据持久化协同关系示意": module.create_architecture_diagram(),
        "自治闭环、知识检索与执行边界协同流程示意": module.create_flow_diagram(),
    }
    for doc_path in DOC_PATHS:
        if not doc_path.exists():
            print(f"skip missing: {doc_path}")
            continue
        for caption, image_path in assets.items():
            refresh_doc(doc_path, caption, image_path)
        print(doc_path)


if __name__ == "__main__":
    main()
