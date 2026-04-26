from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFilter, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


ROOT = Path("/Users/flx/Documents/jienengjianpai")
DOC_PATHS = [
    ROOT / "校园交换机智能体节能平台设计与实现.docx",
    ROOT / "校园交换机智能体节能平台设计与实现-加图版.docx",
]
SHOT_DIR = Path("/tmp/doc-ui-shots/png")
ASSET_DIR = Path("/tmp/jnjp-doc-figures")
ASSET_DIR.mkdir(parents=True, exist_ok=True)

FONT_SANS = "/System/Library/Fonts/Supplemental/Songti.ttc"
BG = "#F6FBFC"
TEXT = "#253847"
MUTED = "#708896"
BORDER = "#D6E4E8"
CARD = "#FFFFFF"
TEAL = "#2E5A64"
PILL = "#EAF5F6"


def load_font(path: str, size: int) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(path, size)


def text_size(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont) -> tuple[int, int]:
    bbox = draw.textbbox((0, 0), text, font=font)
    return bbox[2] - bbox[0], bbox[3] - bbox[1]


def draw_text_center(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font: ImageFont.FreeTypeFont,
    fill: str,
):
    tw, th = text_size(draw, text, font)
    x = box[0] + (box[2] - box[0] - tw) / 2
    y = box[1] + (box[3] - box[1] - th) / 2 - 2
    draw.text((x, y), text, font=font, fill=fill)


def rounded_card(
    img: Image.Image,
    box: tuple[int, int, int, int],
    radius: int,
    fill: str,
    outline: str | None = None,
    shadow_offset: tuple[int, int] = (0, 10),
    shadow_blur: int = 22,
    shadow_color: str = "#D6E4E8",
):
    shadow = Image.new("RGBA", img.size, (0, 0, 0, 0))
    sdraw = ImageDraw.Draw(shadow)
    sx0, sy0, sx1, sy1 = box
    dx, dy = shadow_offset
    sdraw.rounded_rectangle((sx0 + dx, sy0 + dy, sx1 + dx, sy1 + dy), radius=radius, fill=shadow_color)
    shadow = shadow.filter(ImageFilter.GaussianBlur(shadow_blur))
    img.alpha_composite(shadow)
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=2 if outline else 0)


def insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def remove_paragraph(paragraph: Paragraph):
    parent = paragraph._element.getparent()
    if parent is not None:
        parent.remove(paragraph._element)


def paragraph_has_drawing(paragraph: Paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing"))


def paragraph_has_page_break(paragraph: Paragraph) -> bool:
    return bool(paragraph._p.xpath('.//w:br[@w:type="page"]'))


def crop_ui_shot(name: str, box: tuple[int, int, int, int]) -> Image.Image:
    img = Image.open(SHOT_DIR / f"{name}.png").convert("RGB")
    return img.crop(box)


def create_single_ui_panel(label: str, key: str, crop_box: tuple[int, int, int, int]) -> Path:
    canvas = Image.new("RGBA", (920, 1320), BG)
    rounded_card(canvas, (38, 34, 882, 1282), radius=40, fill=CARD, outline="#E4EEF0")
    draw = ImageDraw.Draw(canvas)
    sans_22 = load_font(FONT_SANS, 22)

    draw.rounded_rectangle((96, 90, 382, 142), radius=24, fill=PILL)
    draw_text_center(draw, (108, 94, 370, 140), label, sans_22, TEAL)

    shot = crop_ui_shot(key, crop_box)
    shot = shot.resize((740, 1088), Image.Resampling.LANCZOS)
    canvas.paste(shot, (90, 168))
    draw.rounded_rectangle((90, 168, 830, 1256), radius=28, outline=BORDER, width=3)

    path = ASSET_DIR / f"{key}-single.png"
    canvas.convert("RGB").save(path, quality=95)
    return path


def add_picture_after(paragraph: Paragraph, image_path: Path, width_cm: float) -> Paragraph:
    img_para = insert_paragraph_after(paragraph)
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.paragraph_format.keep_together = True
    img_para.paragraph_format.space_before = Pt(5)
    img_para.paragraph_format.space_after = Pt(10)
    run = img_para.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    return img_para


def add_page_break_after(paragraph: Paragraph) -> Paragraph:
    break_para = insert_paragraph_after(paragraph)
    break_para.paragraph_format.space_before = Pt(0)
    break_para.paragraph_format.space_after = Pt(0)
    break_para.add_run().add_break(WD_BREAK.PAGE)
    return break_para


def build_assets() -> dict[str, Path]:
    crop_boxes = {
        "overview-mid-visible": (14, 80, 476, 760),
        "agent-overview-visible": (14, 80, 476, 760),
        "agent-guardrails-visible": (14, 80, 476, 760),
        "insight-mid-visible": (14, 10, 476, 760),
        "audit-visible": (14, 10, 476, 760),
    }
    panels = {
        "图1系统总览驾驶舱：集中展示系统运行状态、核心收益、处理结果": ("图1 系统总览驾驶舱", "overview-mid-visible"),
        "图2自治智能体主控台：配置自治目标、查看运行状态与累计成果": ("图2 自治主控台", "agent-overview-visible"),
        "图3执行边界（护栏）配置页面：护栏规则配置，保障可控自治": ("图3 执行边界", "agent-guardrails-visible"),
        "图4节能统计与绿色收益可视化：节能成果量化展示，便于汇报": ("图4 统计分析", "insight-mid-visible"),
        "图5审计日志与执行记录页面：全流程操作追溯，满足管理要求": ("图5 审计追踪", "audit-visible"),
    }
    return {
        paragraph_text: create_single_ui_panel(label, key, crop_boxes[key])
        for paragraph_text, (label, key) in panels.items()
    }


def clear_old_gallery(doc: Document):
    paragraphs = list(doc.paragraphs)
    for i, paragraph in enumerate(paragraphs):
        if paragraph.text == "系统五大核心页面组合展示":
            if i > 0 and paragraph_has_drawing(paragraphs[i - 1]):
                remove_paragraph(paragraphs[i - 1])
            remove_paragraph(paragraph)
            break


def clear_old_single_panels(doc: Document, targets: set[str]):
    changed = True
    while changed:
        changed = False
        paragraphs = list(doc.paragraphs)
        for i, paragraph in enumerate(paragraphs[:-1]):
            if paragraph.text in targets and paragraph_has_drawing(paragraphs[i + 1]):
                remove_paragraph(paragraphs[i + 1])
                if i + 2 < len(paragraphs) and paragraph_has_page_break(paragraphs[i + 2]):
                    remove_paragraph(paragraphs[i + 2])
                changed = True
                break


def update_doc(doc_path: Path, assets: dict[str, Path]):
    doc = Document(doc_path)
    clear_old_gallery(doc)
    clear_old_single_panels(doc, set(assets))

    targets = list(assets.keys())
    for paragraph in doc.paragraphs:
        if paragraph.text == "作品实物展示":
            paragraph.paragraph_format.page_break_before = True
            paragraph.paragraph_format.keep_with_next = True

    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text in assets:
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.space_after = Pt(4)
            pic = add_picture_after(paragraph, assets[paragraph.text], 10.8)
            if paragraph.text != targets[-1]:
                add_page_break_after(pic)

    doc.save(doc_path)


def main():
    assets = build_assets()
    for doc_path in DOC_PATHS:
        update_doc(doc_path, assets)
    for doc_path in DOC_PATHS:
        print(doc_path)


if __name__ == "__main__":
    main()
